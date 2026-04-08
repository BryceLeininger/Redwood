"""Tests for the local deal-folder workflow."""

from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path

from docx import Document

from land_due_diligence_agent.config import Settings
from land_due_diligence_agent.deal_pipeline import run_local_deal_pipeline


class LocalDealPipelineTests(unittest.TestCase):
    def test_pipeline_writes_stable_manifest_registry_and_docx_report(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            deal_folder, source_drop = self._build_sample_deal(Path(temp_dir))

            settings = Settings(log_level="INFO")
            result, exit_code = run_local_deal_pipeline(deal_folder, settings=settings)

            self.assertEqual(exit_code, 0)
            self.assertTrue(Path(result.manifest_json_path).samefile(deal_folder / "03_Metadata" / "deal_manifest.json"))
            self.assertTrue(Path(result.issue_registry_path).samefile(deal_folder / "03_Metadata" / "issue_registry.json"))
            self.assertTrue(Path(result.review_report_path).samefile(deal_folder / "04_Output" / "Due_Diligence_Review.docx"))
            self.assertTrue(Path(result.run_log_path).samefile(deal_folder / "04_Output" / "run_log.txt"))
            self.assertEqual(result.latest_run_path, "")
            self.assertTrue(Path(result.manifest_json_path).exists())
            self.assertTrue(Path(result.issue_registry_path).exists())
            self.assertTrue(Path(result.review_report_path).exists())
            self.assertTrue(Path(result.run_log_path).exists())

            manifest_payload = json.loads(Path(result.manifest_json_path).read_text(encoding="utf-8"))
            self.assertEqual(manifest_payload["file_count"], 4)
            statuses = {
                item["file_name"]: item["extraction_status"]
                for item in manifest_payload["files"]
            }
            self.assertEqual(statuses["image.png"], "unsupported")
            self.assertEqual(statuses["purchase_agreement.txt"], "success")

            registry_payload = json.loads(Path(result.issue_registry_path).read_text(encoding="utf-8"))
            self.assertTrue(registry_payload["facts"])
            self.assertTrue(registry_payload["conflicts"])
            self.assertTrue(registry_payload["seller_questions"])

            review_doc = Document(result.review_report_path)
            review_text = "\n".join(paragraph.text for paragraph in review_doc.paragraphs if paragraph.text)
            paragraph_styles = {
                paragraph.text: paragraph.style.name
                for paragraph in review_doc.paragraphs
                if paragraph.text
            }
            self.assertIn("Executive Summary", review_text)
            self.assertIn("Deal Overview", review_text)
            self.assertIn("Entitlement & Zoning", review_text)
            self.assertIn("Site & Product", review_text)
            self.assertIn("Title & Ownership", review_text)
            self.assertIn("Environmental & Geotech", review_text)
            self.assertIn("Utilities & Infrastructure", review_text)
            self.assertIn("Fees / Cost Drivers", review_text)
            self.assertIn("Key Risks & Open Issues", review_text)
            self.assertIn("Missing Information", review_text)
            self.assertIn("Questions for Seller", review_text)
            self.assertIn("Deal Snapshot", review_text)
            self.assertEqual(paragraph_styles["Executive Summary"], "Heading 1")
            self.assertEqual(paragraph_styles["Deal Overview"], "Heading 1")
            self.assertEqual(paragraph_styles["Deal Snapshot"], "Heading 2")
            self.assertNotIn("Top Critical Issues", review_text)
            self.assertNotIn("Detailed Findings by Category", review_text)
            self.assertNotIn("Contradictions / Tensions", review_text)
            self.assertIn("Purchase price referenced at $12,500,000.", review_text)

            extraction_files = [
                path
                for path in (deal_folder / "02_Text_Extraction").rglob("*")
                if path.is_file()
            ]
            self.assertEqual(extraction_files, [])

            source_files = sorted(path.name for path in source_drop.iterdir())
            self.assertEqual(
                source_files,
                [
                    "environment_notes.csv",
                    "image.png",
                    "purchase_agreement.txt",
                    "title_report.txt",
                ],
            )

    def test_debug_mode_writes_timestamped_debug_artifacts(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            deal_folder, _ = self._build_sample_deal(Path(temp_dir))

            settings = Settings(log_level="INFO", debug_mode=True)
            result, exit_code = run_local_deal_pipeline(deal_folder, settings=settings)

            self.assertEqual(exit_code, 0)
            self.assertTrue((deal_folder / "02_Text_Extraction" / result.run_id / "purchase_agreement.txt.txt").exists())
            self.assertTrue((deal_folder / "02_Text_Extraction" / result.run_id / "purchase_agreement.txt.json").exists())
            self.assertTrue((deal_folder / "03_Metadata" / result.run_id / "run_summary.json").exists())
            self.assertTrue((deal_folder / "04_Output" / result.run_id / "Due_Diligence_Review.docx").exists())
            self.assertTrue(Path(result.review_report_path).exists())

    def _build_sample_deal(self, root: Path) -> tuple[Path, Path]:
        deal_folder = root / "d1_375Diana"
        source_drop = deal_folder / "00_Source_Drop"
        (deal_folder / "01_Working").mkdir(parents=True)
        source_drop.mkdir(parents=True)

        (source_drop / "purchase_agreement.txt").write_text(
            "Purchase Price: $12,500,000. APN: 123-456-78. Zoning: R-1. 84 lots.",
            encoding="utf-8",
        )
        (source_drop / "title_report.txt").write_text(
            "Assessor Parcel Number: 123-456-79. Owner: Diana Land Holdings LLC. Easement exception shown in title.",
            encoding="utf-8",
        )
        (source_drop / "environment_notes.csv").write_text(
            "topic,value\nenvironment,recognized environmental condition\nutilities,will serve pending\n",
            encoding="utf-8",
        )
        (source_drop / "image.png").write_text("not supported", encoding="utf-8")

        return deal_folder, source_drop


if __name__ == "__main__":
    unittest.main()