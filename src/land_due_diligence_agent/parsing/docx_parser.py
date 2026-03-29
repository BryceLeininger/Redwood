"""DOCX extraction helpers."""

from __future__ import annotations

from pathlib import Path

from docx import Document


def extract_docx_text(path: Path) -> tuple[str, dict[str, int], list[str], list[dict[str, str | int | None]]]:
    """Extract paragraphs and table content from a Word document."""

    document = Document(str(path))
    warnings: list[str] = []
    parts: list[str] = []
    chunk_records: list[dict[str, str | int | None]] = []

    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    if paragraphs:
        parts.extend(paragraphs)
        chunk_records.extend(
            {
                "chunk_id": f"paragraph-{index:04d}",
                "page_number": None,
                "text": paragraph,
            }
            for index, paragraph in enumerate(paragraphs, start=1)
        )

    for table_index, table in enumerate(document.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            table_text = f"Table {table_index}\n" + "\n".join(rows)
            parts.append(table_text)
            chunk_records.append(
                {
                    "chunk_id": f"table-{table_index:04d}",
                    "page_number": None,
                    "text": table_text,
                }
            )

    if not parts:
        warnings.append("No DOCX text extracted.")

    metadata = {
        "paragraph_count": len(paragraphs),
        "table_count": len(document.tables),
    }
    return "\n\n".join(parts), metadata, warnings, chunk_records
