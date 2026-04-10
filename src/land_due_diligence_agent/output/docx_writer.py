"""DOCX reporting for the local due diligence workflow."""

from __future__ import annotations

from collections import defaultdict
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
import re

from docx import Document
from docx.shared import Pt

from land_due_diligence_agent.analysis.front_end import (
    cost_exposure_band_for_issue,
    deal_impact_magnitude_for_issue,
    deal_impact_mechanism_for_issue,
    deal_impact_summary_issues,
    deal_impact_type_for_issue,
    fixability_classification_for_issue,
    if_wrong_line_for_issue,
    timing_exposure_band_for_issue,
    underwrite_confidence_level,
    underwrite_confidence_limiters,
    underwrite_confidence_reason,
)
from land_due_diligence_agent.deal_models import ConflictRecord, DealRunResult, FactRecord, MissingItem, SourceReference
from land_due_diligence_agent.models import (
    AcquisitionControllingFact,
    AcquisitionCriticalPathStep,
    AcquisitionRiskItem,
    AcquisitionSanityCorrection,
    CanonicalIssue,
    Citation,
    DealSynthesis,
    DocumentAnalysis,
    OmissionAssessment,
)
from land_due_diligence_agent.utils.files import ensure_directory
from land_due_diligence_agent.utils.text import clip_text, unique_preserve_order

_MATERIAL_ISSUE_CATEGORIES = {
    "Title / Access Concerns",
    "Entitlement Status",
    "Environmental Risks",
    "Geotechnical Risks",
    "Flood / Drainage Issues",
    "Utilities / Infrastructure Issues",
    "Offsite Obligations",
    "Fee / Exaction Burden",
    "Budget / Cost Reliability",
    "Schedule Risks",
}
_CONFLICT_TYPE_ORDER = {
    "purchase_price": 0,
    "gross_acreage": 1,
    "net_acreage": 2,
    "site_acreage": 3,
    "lot_count": 4,
    "unit_count": 5,
    "zoning": 6,
    "jurisdiction": 7,
    "owner_name": 8,
    "apn": 9,
}
_PRODUCT_KEYWORDS: tuple[tuple[str, tuple[str, ...]], ...] = (
    ("single-family detached product", ("single family", "single-family", "detached")),
    ("townhome product", ("townhome", "town house")),
    ("multifamily product", ("multifamily", "apartment", "stacked flat")),
    ("industrial product", ("industrial", "warehouse", "distribution")),
    ("commercial product", ("commercial", "retail", "office")),
)
_SECTION_EMPTY_TEXT = {
    "Entitlement & Zoning": "Current zoning, jurisdiction, and approval status are not cleanly established from readable planning support.",
    "Site & Product": "Controlling acreage, yield, and layout support are not cleanly established from the current package.",
    "Title & Ownership": "Title, vesting, and access support are not complete enough to treat ownership and closability as closed.",
    "Environmental & Geotech": "No decision-grade environmental, geotechnical, or drainage support currently closes this lane.",
    "Utilities & Infrastructure": "Utility, frontage, and infrastructure support remain incomplete or only partially established.",
    "Fees / Cost Drivers": "Fee, budget, and site-cost support are not current enough to lock basis with confidence.",
}
_MISSING_STATUS_ORDER = {
    "missing and important": 0,
    "conflicting across documents": 1,
    "stale and potentially unreliable": 2,
    "missing but normally expected": 3,
}
_FRONT_END_PRIORITY = {
    "red flag": 0,
    "conflict / contradiction concern": 1,
    "yellow flag": 2,
    "stale-information concern": 3,
    "document gap": 4,
    "routine item": 5,
}
_ACQUISITION_SEVERITY_PRIORITY = {
    "CRITICAL": 0,
    "HIGH": 1,
    "MODERATE": 2,
    "LOW": 3,
}
_MISSING_INPUT_PRIORITY = {
    "missing and important": 0,
    "stale and potentially unreliable": 1,
    "missing but normally expected": 2,
}
_CONTROL_DOC_BUCKET_ORDER = {
    "must read personally": 0,
    "should skim": 1,
    "safe to rely on agent": 2,
}
_REAL_RISK_FLAGS = {"red flag", "yellow flag", "conflict / contradiction concern"}
_ACQUISITION_BUCKETS = (
    "Primary Deal Driver",
    "Secondary Drivers",
    "Supporting Risks",
    "Noise",
)


@dataclass(slots=True)
class _SectionBullet:
    text: str
    references: list[str] = field(default_factory=list)
    note: str = ""


@dataclass(slots=True)
class _MissingInputBullet:
    label: str
    summary: str
    request: str
    category: str = ""
    status: str = ""
    priority: int = 9
    references: list[str] = field(default_factory=list)
    document_hint: str = ""


def write_due_diligence_report_docx(path: Path, result: DealRunResult) -> Path:
    """Write the primary due diligence report as a Word document."""

    ensure_directory(path.parent)
    document = Document()
    _configure_document(document)

    document.add_heading("Land Due Diligence Memo", level=0)
    document.add_paragraph(f"Deal: {result.deal_name}")
    document.add_paragraph(f"Generated: {datetime.now().astimezone().isoformat(timespec='minutes')}")

    if result.deal_synthesis is None:
        _write_minimal_report(document, result)
    else:
        _write_structured_report(document, result, result.deal_synthesis)

    document.save(path)
    return path


def _configure_document(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.08

    heading_1 = document.styles["Heading 1"]
    heading_1.font.name = "Aptos"
    heading_1.font.size = Pt(14)
    heading_1.paragraph_format.space_before = Pt(16)
    heading_1.paragraph_format.space_after = Pt(6)
    heading_1.paragraph_format.keep_with_next = True

    heading_2 = document.styles["Heading 2"]
    heading_2.font.name = "Aptos"
    heading_2.font.size = Pt(11.5)
    heading_2.paragraph_format.space_before = Pt(10)
    heading_2.paragraph_format.space_after = Pt(3)
    heading_2.paragraph_format.keep_with_next = True

    list_bullet = document.styles["List Bullet"]
    list_bullet.font.name = "Aptos"
    list_bullet.font.size = Pt(10.5)
    list_bullet.paragraph_format.space_after = Pt(2)

    list_number = document.styles["List Number"]
    list_number.font.name = "Aptos"
    list_number.font.size = Pt(10.5)
    list_number.paragraph_format.space_after = Pt(2)


def _write_minimal_report(document: Document, result: DealRunResult) -> None:
    _add_section(document, "BOTTOM LINE")
    _add_bullet_items(
        document,
        [
            _SectionBullet(
                text="No readable document set was available for decision-grade analysis, so the deal facts, risk profile, and package quality remain unresolved.",
            )
        ],
    )

    _add_section(document, "Missing Information")
    if result.failed_files:
        _add_bullet_items(
            document,
            [
                _SectionBullet(
                    text=f"{result.failed_files} file(s) failed extraction and require direct manual review before relying on the package.",
                )
            ],
        )
    else:
        _add_bullet_items(
            document,
            [
                _SectionBullet(
                    text="The report could not assemble decision-grade support from the current package.",
                )
            ],
        )


def _write_structured_report(document: Document, result: DealRunResult, synthesis: DealSynthesis) -> None:
    fact_index = _build_fact_index(result.issue_registry.facts)
    material_issues = _material_issues(synthesis)
    material_conflicts = _material_conflicts(result.issue_registry.conflicts)
    missing_inputs = _material_missing_inputs(result, synthesis)

    _add_bottom_line_section(
        document=document,
        result=result,
        synthesis=synthesis,
        fact_index=fact_index,
        material_issues=material_issues,
        missing_inputs=missing_inputs,
    )
    _add_deal_impact_summary_section(document, material_issues)
    _add_underwrite_confidence_section(document, synthesis, material_issues)
    _add_sanity_corrections_section(document, synthesis)
    _add_controlling_facts_section(document, synthesis)
    _add_real_risk_classification_section(document, synthesis)
    _add_critical_path_chain_section(document, synthesis)
    _add_investment_decision_section(document, synthesis)
    _add_weak_acquisitions_section(document, synthesis)
    _add_key_documents_that_control_section(document, synthesis)
    _add_deal_overview(document, result, synthesis, fact_index)
    _add_gating_items_section(document, material_issues)
    _add_category_section(
        document=document,
        title="Entitlement & Zoning",
        fact_index=fact_index,
        material_conflicts=material_conflicts,
        material_issues=material_issues,
        missing_inputs=missing_inputs,
        fact_types=("jurisdiction", "zoning"),
        issue_categories={"Entitlement Status"},
        omission_categories={"Entitlement Status"},
        first_pass_missing_categories={"Entitlement / Planning / Conditions"},
        conflict_types={"zoning", "jurisdiction"},
    )
    _add_category_section(
        document=document,
        title="Site & Product",
        fact_index=fact_index,
        material_conflicts=material_conflicts,
        material_issues=material_issues,
        missing_inputs=missing_inputs,
        fact_types=("gross_acreage", "net_acreage", "site_acreage", "lot_count", "unit_count"),
        issue_categories=set(),
        omission_categories=set(),
        first_pass_missing_categories={"Map / Plat / Improvement Plans", "Financial / underwriting support"},
        conflict_types={"gross_acreage", "net_acreage", "site_acreage", "lot_count", "unit_count"},
    )
    _add_category_section(
        document=document,
        title="Title & Ownership",
        fact_index=fact_index,
        material_conflicts=material_conflicts,
        material_issues=material_issues,
        missing_inputs=missing_inputs,
        fact_types=("apn", "owner_name"),
        issue_categories={"Title / Access Concerns"},
        omission_categories={"Title / Access Concerns"},
        first_pass_missing_categories={"Title", "Vesting / Legal"},
        conflict_types={"apn", "owner_name"},
    )
    _add_category_section(
        document=document,
        title="Environmental & Geotech",
        fact_index=fact_index,
        material_conflicts=material_conflicts,
        material_issues=material_issues,
        missing_inputs=missing_inputs,
        fact_types=(),
        issue_categories={"Environmental Risks", "Geotechnical Risks", "Flood / Drainage Issues"},
        omission_categories={"Environmental Risks", "Geotechnical Risks", "Flood / Drainage Issues"},
        first_pass_missing_categories={"Environmental", "Geotech / Soils"},
        conflict_types=set(),
    )
    _add_category_section(
        document=document,
        title="Utilities & Infrastructure",
        fact_index=fact_index,
        material_conflicts=material_conflicts,
        material_issues=material_issues,
        missing_inputs=missing_inputs,
        fact_types=(),
        issue_categories={"Utilities / Infrastructure Issues", "Offsite Obligations"},
        omission_categories={"Utilities / Infrastructure Issues"},
        first_pass_missing_categories={"Utilities", "Map / Plat / Improvement Plans"},
        conflict_types=set(),
    )
    _add_category_section(
        document=document,
        title="Fees / Cost Drivers",
        fact_index=fact_index,
        material_conflicts=material_conflicts,
        material_issues=material_issues,
        missing_inputs=missing_inputs,
        fact_types=("purchase_price",),
        issue_categories={"Fee / Exaction Burden", "Budget / Cost Reliability"},
        omission_categories={"Fee / Exaction Burden", "Budget / Cost Reliability"},
        first_pass_missing_categories={"Purchase / Sale / Contract"},
        conflict_types={"purchase_price"},
    )
    _add_key_risks_section(document, material_issues)
    _add_missing_information_section(document, result, synthesis, missing_inputs)
    _add_questions_for_seller_section(document, material_issues, missing_inputs)


def _add_bottom_line_section(
    *,
    document: Document,
    result: DealRunResult,
    synthesis: DealSynthesis,
    fact_index: dict[str, list[FactRecord]],
    material_issues: list[CanonicalIssue],
    missing_inputs: list[_MissingInputBullet],
) -> None:
    _add_section(document, "BOTTOM LINE")

    _add_subsection(document, "Deal")
    _add_bullet_items(document, [_SectionBullet(text=_bottom_line_deal_text(result, synthesis, fact_index))])

    _add_subsection(document, "Biggest Current Concerns")
    concern_bullets = [
        _SectionBullet(
            text=_bottom_line_issue_text(issue),
            references=_issue_reference_labels(issue),
        )
        for issue in material_issues[:3]
    ]
    _add_bullet_items(
        document,
        concern_bullets or [_SectionBullet(text="No material issue currently rises above routine diligence noise based on the provided documents.")],
    )

    _add_subsection(document, "Most Important Missing Items / Confirmations")
    missing_bullets = [
        _SectionBullet(
            text=_bottom_line_missing_text(item),
            note=f"Best document: {item.document_hint}." if item.document_hint else "",
            references=item.references,
        )
        for item in missing_inputs[:3]
    ]
    _add_bullet_items(
        document,
        missing_bullets or [_SectionBullet(text="No additional critical missing confirmation was isolated beyond the current issue set.")],
    )

    _add_subsection(document, "Package Read")
    _add_bullet_items(
        document,
        [
            _SectionBullet(
                text=_bottom_line_package_read(synthesis, material_issues, missing_inputs),
            )
        ],
    )


def _add_key_documents_that_control_section(document: Document, synthesis: DealSynthesis) -> None:
    _add_section(document, "Key Documents That Control")
    control_documents = _key_documents_that_control(synthesis.document_analyses)
    _add_bullet_items(
        document,
        control_documents or [_SectionBullet(text="The package does not contain an obvious current controlling document set in the lanes that normally anchor deal judgment.")],
    )


def _add_deal_impact_summary_section(document: Document, material_issues: list[CanonicalIssue]) -> None:
    _add_section(document, "DEAL IMPACT SUMMARY")
    summary_issues = deal_impact_summary_issues(material_issues, limit=3)
    bullets = [
        _SectionBullet(
            text=(
                f"{issue.title} [{issue.acquisition_severity}]: impact type {deal_impact_type_for_issue(issue)}; "
                f"magnitude {deal_impact_magnitude_for_issue(issue)}. "
                f"{clip_text(deal_impact_mechanism_for_issue(issue), 180)}"
            ),
            note=(
                f"Cost exposure: {cost_exposure_band_for_issue(issue)}. "
                f"Timing exposure: {timing_exposure_band_for_issue(issue)}. "
                f"Fixability: {fixability_classification_for_issue(issue)}."
            ),
            references=_issue_reference_labels(issue),
        )
        for issue in summary_issues
    ]
    _add_bullet_items(
        document,
        bullets
        or [
            _SectionBullet(
                text="No CRITICAL or HIGH issue currently rises high enough to reshape the deal beyond routine diligence judgment."
            )
        ],
    )


def _add_underwrite_confidence_section(
    document: Document,
    synthesis: DealSynthesis,
    material_issues: list[CanonicalIssue],
) -> None:
    _add_section(document, "UNDERWRITE CONFIDENCE")

    _add_subsection(document, "Current Read")
    confidence_level = underwrite_confidence_level(
        registry=synthesis.canonical_issue_registry,
        omission_assessments=synthesis.omission_assessments,
        contradictions=synthesis.contradictions,
        document_analyses=synthesis.document_analyses,
        issues=material_issues,
    )
    confidence_reason = underwrite_confidence_reason(
        registry=synthesis.canonical_issue_registry,
        omission_assessments=synthesis.omission_assessments,
        contradictions=synthesis.contradictions,
        document_analyses=synthesis.document_analyses,
        issues=material_issues,
    )
    _add_bullet_items(
        document,
        [
            _SectionBullet(
                text=f"Current underwrite confidence is {confidence_level}. {clip_text(confidence_reason, 220)}"
            )
        ],
    )

    _add_subsection(document, "Main Limiters")
    limiter_bullets = [
        _SectionBullet(text=clip_text(line, 200))
        for line in underwrite_confidence_limiters(
            registry=synthesis.canonical_issue_registry,
            omission_assessments=synthesis.omission_assessments,
            contradictions=synthesis.contradictions,
            document_analyses=synthesis.document_analyses,
            issues=material_issues,
            limit=3,
        )
    ]
    _add_bullet_items(
        document,
        limiter_bullets
        or [
            _SectionBullet(
                text="No single unresolved assumption stands out beyond the current issue list and package-read summary."
            )
        ],
    )

    _add_subsection(document, "If Wrong, What Happens?")
    downside_bullets = [
        _SectionBullet(
            text=f"{issue.title}: {clip_text(if_wrong_line_for_issue(issue), 200)}",
            references=_issue_reference_labels(issue),
        )
        for issue in deal_impact_summary_issues(material_issues, limit=3)
    ]
    _add_bullet_items(
        document,
        downside_bullets
        or [
            _SectionBullet(
                text="If the remaining assumptions break, the deal can still move on price, timing, or closability faster than the current package implies."
            )
        ],
    )


def _add_controlling_facts_section(document: Document, synthesis: DealSynthesis) -> None:
    _add_section(document, "Controlling Facts")
    bullets = [_controlling_fact_bullet(fact) for fact in synthesis.acquisition_judgment.controlling_facts]
    _add_bullet_items(
        document,
        bullets
        or [
            _SectionBullet(
                text="The second pass did not isolate a controlling answer for the core underwriting descriptors from the current readable package."
            )
        ],
    )


def _add_sanity_corrections_section(document: Document, synthesis: DealSynthesis) -> None:
    _add_section(document, "Sanity Check / Corrections")
    bullets = [_sanity_correction_bullet(item) for item in synthesis.acquisition_judgment.sanity_corrections]
    _add_bullet_items(
        document,
        bullets
        or [
            _SectionBullet(
                text="No controlling fact required a second-pass sanity correction beyond the current readable package."
            )
        ],
    )


def _add_real_risk_classification_section(document: Document, synthesis: DealSynthesis) -> None:
    _add_section(document, "Real Risk Classification")
    for bucket in _ACQUISITION_BUCKETS:
        _add_subsection(document, bucket)
        items = [item for item in synthesis.acquisition_judgment.risk_items if item.bucket == bucket]
        bullets = [_acquisition_risk_bullet(item) for item in items[:5]]
        _add_bullet_items(
            document,
            bullets or [_SectionBullet(text=f"No item currently lands in {bucket.lower()}.")],
        )


def _add_critical_path_chain_section(document: Document, synthesis: DealSynthesis) -> None:
    _add_section(document, "Critical Path / Gating Chain")
    for target in ("Final Map", "Grading Permit", "Vertical Start"):
        _add_subsection(document, target)
        steps = [step for step in synthesis.acquisition_judgment.critical_path if step.target == target]
        bullets = [_critical_path_bullet(step) for step in steps]
        _add_bullet_items(
            document,
            bullets or [_SectionBullet(text=f"No stage-specific blocker was isolated for {target.lower()} beyond the current general issue set.")],
        )


def _add_investment_decision_section(document: Document, synthesis: DealSynthesis) -> None:
    decision = synthesis.acquisition_judgment.investment_decision
    _add_section(document, "Investment Decision")

    _add_subsection(document, "Primary Driver")
    _add_bullet_items(
        document,
        [_SectionBullet(text=decision.primary_driver)]
        or [_SectionBullet(text="No single primary deal driver is isolated beyond routine execution noise.")],
    )

    _add_subsection(document, "Secondary Drivers")
    _add_bullet_items(
        document,
        [_SectionBullet(text=line) for line in decision.secondary_drivers]
        or [_SectionBullet(text="No secondary driver rises above supporting execution risk.")],
    )

    _add_subsection(document, "Decision")
    _add_bullet_items(
        document,
        [
            _SectionBullet(
                text=f"{decision.posture}. {decision.rationale}",
                references=_citation_labels(decision.citations),
            )
        ],
    )

    _add_subsection(document, "Top 3 Real Risks")
    _add_bullet_items(
        document,
        [_SectionBullet(text=line) for line in decision.top_real_risks]
        or [_SectionBullet(text="No real risk currently rises above routine diligence noise in the second-pass classification.")],
    )

    _add_subsection(document, "Price / Structure Changes")
    _add_bullet_items(
        document,
        [_SectionBullet(text=line) for line in decision.price_or_structure_changes]
        or [_SectionBullet(text="No specific price or structure change currently rises above routine contingency management.")],
    )

    _add_subsection(document, "Single Biggest Unknown")
    _add_bullet_items(
        document,
        [
            _SectionBullet(
                text=decision.biggest_unknown or "No single unknown currently stands above the rest of the issue set.",
                references=_citation_labels(decision.citations),
            )
        ],
    )

    _add_subsection(document, "What Has To Be True")
    _add_bullet_items(
        document,
        [_SectionBullet(text=line) for line in decision.what_has_to_be_true]
        or [_SectionBullet(text="No additional gating condition rises above the current reset risk ranking.")],
    )

    _add_subsection(document, "What Must Be True To Close")
    _add_bullet_items(
        document,
        [_SectionBullet(text=line) for line in decision.close_requirements]
        or [_SectionBullet(text="No close blocker is currently isolated beyond routine closing work.")],
    )

    _add_subsection(document, "What Must Be True To Start Grading")
    _add_bullet_items(
        document,
        [_SectionBullet(text=line) for line in decision.grading_requirements]
        or [_SectionBullet(text="No grading-start blocker is currently isolated beyond routine engineering coordination.")],
    )

    _add_subsection(document, "What Must Be True For Vertical Start")
    _add_bullet_items(
        document,
        [_SectionBullet(text=line) for line in decision.vertical_requirements]
        or [_SectionBullet(text="No vertical-start blocker is currently isolated beyond routine budget and permit coordination.")],
    )

    _add_subsection(document, "Risks Underwritten")
    _add_bullet_items(
        document,
        [_SectionBullet(text=line) for line in decision.risks_underwritten]
        or [_SectionBullet(text="No execution risk is currently being affirmatively underwritten beyond routine project friction.")],
    )

    _add_subsection(document, "Treated As Solved")
    _add_bullet_items(
        document,
        [_SectionBullet(text=line) for line in decision.treated_as_solved]
        or [_SectionBullet(text="No lane should currently be treated as solved beyond document-backed descriptors.")],
    )


def _add_weak_acquisitions_section(document: Document, synthesis: DealSynthesis) -> None:
    _add_section(document, "What A Weak Acquisitions Person Would Miss")
    bullets = [
        _SectionBullet(
            text=f"{insight.title}: {insight.detail}",
            references=[*_citation_labels(insight.citations), *insight.source_documents],
        )
        for insight in synthesis.acquisition_judgment.weak_acquisition_misses
    ]
    _add_bullet_items(
        document,
        bullets or [_SectionBullet(text="The second pass did not isolate three non-obvious points beyond the existing issue ranking.")],
    )


def _add_executive_summary(
    *,
    document: Document,
    result: DealRunResult,
    synthesis: DealSynthesis,
    fact_index: dict[str, list[FactRecord]],
    material_issues: list[CanonicalIssue],
    critical_missing: list[OmissionAssessment],
) -> None:
    _add_section(document, "Executive Summary")

    _add_subsection(document, "IC Read")
    _add_bullet_items(
        document,
        [
            _SectionBullet(text=clip_text(synthesis.executive_summary, 420)),
            _SectionBullet(
                text=(
                    f"Current recommendation posture: {synthesis.recommendation.posture}. "
                    f"This is being driven by {', '.join(issue.title for issue in material_issues[:2]) or 'the current package quality and unresolved diligence items'}.")
            ),
        ],
    )

    _add_subsection(document, "Deal Snapshot")
    _add_bullet_items(
        document,
        [
            _build_snapshot_bullet(result, synthesis, fact_index),
            _SectionBullet(
                text=(
                    f"Package quality currently reads as {synthesis.canonical_issue_registry.package_quality or 'mixed'} "
                    f"with {synthesis.canonical_issue_registry.confidence_in_initial_read} confidence on initial read."
                ),
            ),
        ],
    )

    _add_subsection(document, "Known With High Confidence")
    high_confidence = _build_high_confidence_bullets(fact_index, material_issues)
    if not high_confidence:
        high_confidence = [
            _SectionBullet(
                text="No core deal descriptor appears clearly in multiple readable documents; location, scale, and product should still be treated as provisional.",
            )
        ]
    _add_bullet_items(document, high_confidence)

    _add_subsection(document, "Top Risks")
    top_risks = [
        _SectionBullet(
            text=f"{issue.title}: {_issue_summary_line(issue)}",
            references=_issue_reference_labels(issue),
        )
        for issue in material_issues[:5]
    ]
    _add_bullet_items(
        document,
        top_risks or [_SectionBullet(text="No material risk was isolated above routine diligence noise in the current package.")],
    )

    _add_subsection(document, "Missing Critical Inputs")
    missing_summary = _missing_summary_bullets(result, critical_missing)
    _add_bullet_items(
        document,
        missing_summary or [_SectionBullet(text="No additional critical missing item was isolated beyond the current issue set.")],
    )


def _add_gating_items_section(
    document: Document,
    material_issues: list[CanonicalIssue],
) -> None:
    _add_section(document, "Deal Killers / Gating Items")
    gating_issues = [
        issue
        for issue in material_issues
        if issue.gating_item or (issue.blocking_flag and issue.acquisition_severity in {"CRITICAL", "HIGH"})
    ][:4]
    if not gating_issues:
        _add_bullet_items(
            document,
            [_SectionBullet(text="No single deal killer is confirmed from the provided documents, but trust in the deal still depends on clearing the lead high-severity items.")],
        )
        return

    _add_bullet_items(
        document,
        [
            _SectionBullet(
                text=f"{issue.title} [{issue.acquisition_severity}]: {_issue_request_text(issue)}",
                note=f"Affects {', '.join(issue.affects) or 'deal execution'}. {_issue_deal_impact(issue)}",
                references=_issue_reference_labels(issue),
            )
            for issue in gating_issues
        ],
    )


def _add_deal_overview(
    document: Document,
    result: DealRunResult,
    synthesis: DealSynthesis,
    fact_index: dict[str, list[FactRecord]],
) -> None:
    _add_section(document, "Deal Overview")
    bullets = [
        _build_location_bullet(fact_index),
        _build_scale_bullet(fact_index),
        _build_product_bullet(result, fact_index),
        _build_price_bullet(fact_index),
        _SectionBullet(
            text=f"Entitlement stage: {clip_text(synthesis.entitlement_status, 180)}",
        ),
    ]
    _add_bullet_items(document, [bullet for bullet in bullets if bullet.text])


def _add_category_section(
    *,
    document: Document,
    title: str,
    fact_index: dict[str, list[FactRecord]],
    material_conflicts: list[ConflictRecord],
    material_issues: list[CanonicalIssue],
    missing_inputs: list[_MissingInputBullet],
    fact_types: tuple[str, ...],
    issue_categories: set[str],
    omission_categories: set[str],
    first_pass_missing_categories: set[str],
    conflict_types: set[str],
) -> None:
    _add_section(document, title)
    bullets = _build_category_bullets(
        fact_index=fact_index,
        material_conflicts=material_conflicts,
        material_issues=material_issues,
        missing_inputs=missing_inputs,
        fact_types=fact_types,
        issue_categories=issue_categories,
        omission_categories=omission_categories,
        first_pass_missing_categories=first_pass_missing_categories,
        conflict_types=conflict_types,
    )
    _add_bullet_items(
        document,
        bullets or [_SectionBullet(text=_SECTION_EMPTY_TEXT[title])],
    )


def _add_key_risks_section(document: Document, material_issues: list[CanonicalIssue]) -> None:
    _add_section(document, "Real Risks / Open Issues")
    if not material_issues:
        _add_bullet_items(
            document,
            [_SectionBullet(text="No material issue rose above routine diligence noise in the current package.")],
        )
        return

    for issue in material_issues[:5]:
        _add_subsection(document, issue.title)
        _add_bullet_items(
            document,
            [
                _SectionBullet(text=f"Severity: {_issue_severity_line(issue)}"),
                _SectionBullet(text=f"Issue: {_issue_what_line(issue)}", references=_issue_reference_labels(issue)),
                _SectionBullet(text=f"Likely explanation: {_issue_likely_explanation(issue)}"),
                _SectionBullet(
                    text=(
                        f"Deal impact: type {deal_impact_type_for_issue(issue)}; "
                        f"magnitude {deal_impact_magnitude_for_issue(issue)}; "
                        f"mechanism {clip_text(deal_impact_mechanism_for_issue(issue), 170)}"
                    )
                ),
                _SectionBullet(
                    text=(
                        f"Exposure / fixability: cost {cost_exposure_band_for_issue(issue)}; "
                        f"timing {timing_exposure_band_for_issue(issue)}; "
                        f"fixability {fixability_classification_for_issue(issue)}."
                    )
                ),
                _SectionBullet(text=f"If wrong: {clip_text(if_wrong_line_for_issue(issue), 200)}"),
                _SectionBullet(text=f"What would resolve it: {_issue_request_text(issue)}"),
            ],
        )


def _add_recommended_next_steps_section(
    document: Document,
    synthesis: DealSynthesis,
    material_issues: list[CanonicalIssue],
    critical_missing: list[OmissionAssessment],
) -> None:
    _add_section(document, "Recommended Next Steps")
    roadmap_steps = synthesis.further_diligence_roadmap.recommended_next_steps[:8]
    if roadmap_steps:
        for step in roadmap_steps:
            document.add_paragraph(step, style="List Number")
        return

    fallback_steps = _build_seller_request_items(material_issues, critical_missing)[:8]
    if not fallback_steps:
        _add_bullet_items(document, [_SectionBullet(text="No additional concrete next step was isolated from the current package.")])
        return

    for item in fallback_steps:
        document.add_paragraph(item.text, style="List Number")
        if item.note:
            _add_note_line(document, f"Why: {item.note}")
        if item.references:
            _add_reference_line(document, item.references)


def _add_missing_information_section(
    document: Document,
    result: DealRunResult,
    synthesis: DealSynthesis,
    missing_inputs: list[_MissingInputBullet],
) -> None:
    _add_section(document, "Missing Information")
    bullets = [
        _SectionBullet(
            text=f"{item.label}: {clip_text(item.summary.rstrip('.') + '.', 180)}",
            note=(f"Request: {item.request.rstrip('.')}." if item.request else "") + (f" Best document: {item.document_hint}." if item.document_hint else ""),
            references=item.references,
        )
        for item in missing_inputs
    ]

    if result.failed_files:
        bullets.append(
            _SectionBullet(
                text=f"{result.failed_files} file(s) failed extraction, so manual review is still required before treating the package as complete.",
            )
        )

    if synthesis.extraction_errors:
        bullets.append(
            _SectionBullet(
                text="Some source files produced extraction errors; any conclusion that depends on those files should be confirmed manually.",
            )
        )

    _add_bullet_items(
        document,
        bullets or [_SectionBullet(text="No additional missing-information item was isolated from the current package.")],
    )


def _add_questions_for_seller_section(
    document: Document,
    material_issues: list[CanonicalIssue],
    missing_inputs: list[_MissingInputBullet],
) -> None:
    _add_section(document, "Questions for Seller")
    questions = _build_seller_request_items(material_issues, missing_inputs)
    if not questions:
        _add_bullet_items(
            document,
            [_SectionBullet(text="No additional seller follow-up question was isolated from the current package.")],
        )
        return

    for item in questions:
        document.add_paragraph(item.text, style="List Number")
        if item.note:
            _add_note_line(document, item.note)


def _build_snapshot_bullet(
    result: DealRunResult,
    synthesis: DealSynthesis,
    fact_index: dict[str, list[FactRecord]],
) -> _SectionBullet:
    location = _location_text(fact_index)
    scale = _scale_text(fact_index)
    product = _product_text(result, fact_index)
    entitlement = clip_text(synthesis.entitlement_status, 140)
    text = f"{result.deal_name}: {location}; {scale}; {product}; entitlement stage currently reads as {entitlement}."

    references: list[str] = []
    for fact_type in ("jurisdiction", "gross_acreage", "net_acreage", "site_acreage", "lot_count", "unit_count"):
        references.extend(_reference_labels_for_fact_bundle(fact_index, fact_type))

    return _SectionBullet(text=text, references=references[:3])


def _build_high_confidence_bullets(
    fact_index: dict[str, list[FactRecord]],
    material_issues: list[CanonicalIssue],
) -> list[_SectionBullet]:
    bullets: list[_SectionBullet] = []
    for fact_type in (
        "jurisdiction",
        "zoning",
        "purchase_price",
        "gross_acreage",
        "net_acreage",
        "site_acreage",
        "lot_count",
        "unit_count",
        "owner_name",
        "apn",
    ):
        fact, supporting_facts = _best_fact_bundle(fact_index, fact_type, require_high=True)
        if fact is None:
            continue
        bullets.append(
            _SectionBullet(
                text=f"{_format_fact_sentence(fact)} This appears consistently across multiple readable documents.",
                references=_reference_labels_from_facts(supporting_facts),
            )
        )
        if len(bullets) >= 3:
            break

    for issue in material_issues:
        if issue.confidence != "high" or issue.information_status != "present and adequate":
            continue
        bullets.append(
            _SectionBullet(
                text=f"{issue.title}: {_issue_known_line(issue)}",
                references=_issue_reference_labels(issue),
            )
        )
        if len(bullets) >= 4:
            break

    return _dedupe_bullets(bullets)[:4]


def _missing_summary_bullets(
    result: DealRunResult,
    critical_missing: list[OmissionAssessment],
) -> list[_SectionBullet]:
    bullets = [
        _SectionBullet(
            text=_omission_text(assessment),
            references=_omission_reference_labels(assessment),
        )
        for assessment in critical_missing[:4]
    ]

    first_pass_items = [
        _SectionBullet(text=_missing_item_text(item))
        for item in result.issue_registry.missing_items[:4]
    ]
    return _dedupe_bullets([*bullets, *first_pass_items])[:5]


def _build_category_bullets(
    *,
    fact_index: dict[str, list[FactRecord]],
    material_conflicts: list[ConflictRecord],
    material_issues: list[CanonicalIssue],
    missing_inputs: list[_MissingInputBullet],
    fact_types: tuple[str, ...],
    issue_categories: set[str],
    omission_categories: set[str],
    first_pass_missing_categories: set[str],
    conflict_types: set[str],
) -> list[_SectionBullet]:
    bullets: list[_SectionBullet] = []

    for fact_type in fact_types:
        fact, supporting_facts = _best_fact_bundle(fact_index, fact_type)
        if fact is None:
            continue
        bullets.append(
            _SectionBullet(
                text=f"Current read: {_fact_section_text(fact, supporting_facts)}",
                references=_reference_labels_from_facts(supporting_facts),
            )
        )

    issues = [issue for issue in material_issues if issue.category in issue_categories]
    for issue in issues[:2]:
        bullets.append(
            _SectionBullet(
                text=f"Open risk: {issue.title}. {_issue_section_detail(issue)}",
                references=_issue_reference_labels(issue),
                note=f"To clear: {_issue_request_text(issue)}",
            )
        )

    for conflict in material_conflicts:
        if conflict.fact_type not in conflict_types:
            continue
        bullets.append(
            _SectionBullet(
                text=f"Control point to resolve: {conflict.description}",
                references=_source_reference_labels(conflict.sources),
                note=conflict.uncertainty,
            )
        )

    omissions = [
        item
        for item in missing_inputs
        if item.category in omission_categories or item.category in first_pass_missing_categories
    ]
    for item in omissions[:2]:
        bullets.append(
            _SectionBullet(
                text=f"Still needed: {item.label}. {clip_text(item.summary.rstrip('.') + '.', 160)}",
                references=item.references,
                note=(f"Request: {item.request.rstrip('.')}." if item.request else "") + (f" Best document: {item.document_hint}." if item.document_hint else ""),
            )
        )

    return _dedupe_bullets(bullets)[:5]


def _build_seller_request_items(
    material_issues: list[CanonicalIssue],
    missing_inputs: list[_MissingInputBullet],
) -> list[_SectionBullet]:
    items: list[_SectionBullet] = []

    ordered_issues = sorted(
        material_issues,
        key=lambda issue: (
            _ACQUISITION_SEVERITY_PRIORITY.get(issue.acquisition_severity, 4),
            0 if issue.gating_item else 1,
            -issue.priority_score.total,
            issue.title,
        ),
    )

    for issue in ordered_issues[:4]:
        items.append(
            _SectionBullet(
                text=_issue_seller_question(issue),
                note=(f"Document that should answer this: {_issue_document_hint(issue)}." if _issue_document_hint(issue) else ""),
            )
        )

    for item in missing_inputs[:4]:
        items.append(
            _SectionBullet(
                text=_missing_input_seller_question(item),
                note=(f"Document that should answer this: {item.document_hint}." if item.document_hint else ""),
            )
        )

    return _dedupe_bullets(items)[:8]


def _build_fact_index(facts: list[FactRecord]) -> dict[str, list[FactRecord]]:
    fact_index: dict[str, list[FactRecord]] = defaultdict(list)
    for fact in facts:
        if fact.fact_type.startswith("signal_") or fact.confidence == "low":
            continue
        fact_index[fact.fact_type].append(fact)
    return fact_index


def _best_fact_bundle(
    fact_index: dict[str, list[FactRecord]],
    fact_type: str,
    *,
    require_high: bool = False,
) -> tuple[FactRecord | None, list[FactRecord]]:
    candidates = fact_index.get(fact_type, [])
    if require_high:
        candidates = [fact for fact in candidates if fact.confidence == "high"]
    if not candidates:
        return None, []

    best = max(
        candidates,
        key=lambda fact: (
            _confidence_rank(fact.confidence),
            _support_count(candidates, fact.normalized_value),
            len(fact.sources),
        ),
    )
    supporting_facts = [fact for fact in candidates if fact.normalized_value == best.normalized_value]
    return best, supporting_facts


def _support_count(facts: list[FactRecord], normalized_value: str) -> int:
    paths: set[str] = set()
    for fact in facts:
        if fact.normalized_value != normalized_value:
            continue
        for source in fact.sources:
            paths.add(source.relative_path)
    return len(paths)


def _material_issues(synthesis: DealSynthesis) -> list[CanonicalIssue]:
    issues = _dedupe_issues(
        [issue for issue in synthesis.canonical_issue_registry.issues if _issue_is_material(issue)]
    )
    issues.sort(
        key=lambda issue: (
            _ACQUISITION_SEVERITY_PRIORITY.get(issue.acquisition_severity, 4),
            _FRONT_END_PRIORITY.get(issue.front_end_flag, 9),
            0 if issue.gating_item else 1,
            0 if issue.blocking_flag else 1,
            0 if issue.critical_path_flag else 1,
            -issue.priority_score.total,
            issue.title,
        )
    )
    return issues[:5]


def _issue_is_material(issue: CanonicalIssue) -> bool:
    if issue.category not in _MATERIAL_ISSUE_CATEGORIES:
        return False
    if issue.front_end_flag == "routine item":
        return False
    if issue.front_end_flag in {"document gap", "stale-information concern"} and issue.information_status != "conflicting across documents":
        return False
    if issue.evidence_basis in {"omission_only", "routine_missing_support"} and issue.information_status != "conflicting across documents":
        return False
    if issue.specificity_level == "generic" and not issue.site_specific_trigger and issue.information_status != "conflicting across documents" and not issue.blocking_flag:
        return False
    if issue.false_positive_risk == "high" and not (issue.gating_item or issue.blocking_flag or issue.acquisition_severity in {"CRITICAL", "HIGH"}):
        return False
    if not _issue_has_clean_support(issue) and issue.information_status != "conflicting across documents" and not issue.blocking_flag:
        return False
    if issue.blocking_flag or issue.critical_path_flag:
        return True
    if issue.acquisition_severity in {"CRITICAL", "HIGH"}:
        return True
    if issue.front_end_flag in {"red flag", "conflict / contradiction concern"}:
        return True
    if issue.front_end_flag == "yellow flag" and _issue_priority_signal(issue) >= 4:
        return True
    return issue.top_line_eligible and _issue_priority_signal(issue) >= 4


def _issue_priority_signal(issue: CanonicalIssue) -> int:
    return max(
        issue.priority_score.cost_exposure,
        issue.priority_score.schedule_exposure,
        issue.priority_score.entitlement_fragility,
        issue.priority_score.closing_risk,
        issue.priority_score.yield_exposure,
    )


def _issue_has_clean_support(issue: CanonicalIssue) -> bool:
    candidates = [
        issue.site_specific_trigger,
        *issue.best_evidence[:2],
        *issue.core_facts[:2],
        issue.why_it_matters,
        issue.likely_implication,
    ]
    return any(candidate and not _text_reads_like_extraction_noise(candidate) for candidate in candidates)


def _text_reads_like_extraction_noise(text: str) -> bool:
    normalized = " ".join(text.split())
    if not normalized:
        return True
    tokens = re.findall(r"[A-Za-z0-9$%/-]+", normalized)
    if len(tokens) < 4:
        return True
    alpha_tokens = [token for token in tokens if re.search(r"[A-Za-z]", token)]
    if not alpha_tokens:
        return True

    cleaned_alpha = [re.sub(r"[^A-Za-z]", "", token) for token in alpha_tokens]
    short_ratio = sum(len(token) <= 2 for token in cleaned_alpha if token) / max(len(alpha_tokens), 1)
    upper_ratio = sum(token.isupper() and len(token) <= 4 for token in alpha_tokens) / max(len(alpha_tokens), 1)
    digit_ratio = sum(character.isdigit() for character in normalized) / len(normalized)
    symbol_ratio = sum(not character.isalnum() and not character.isspace() for character in normalized) / len(normalized)
    frequencies = defaultdict(int)
    for token in (token.lower() for token in alpha_tokens):
        frequencies[token] += 1
    repeated_ratio = max(frequencies.values(), default=0) / max(len(alpha_tokens), 1)

    return (
        (upper_ratio >= 0.55 and short_ratio >= 0.45)
        or digit_ratio >= 0.35
        or symbol_ratio >= 0.18
        or repeated_ratio >= 0.35
    )


def _dedupe_issues(issues: list[CanonicalIssue]) -> list[CanonicalIssue]:
    seen: set[tuple[str, str]] = set()
    deduped: list[CanonicalIssue] = []
    for issue in issues:
        key = (issue.category.lower(), issue.title.strip().lower())
        if key in seen:
            continue
        seen.add(key)
        deduped.append(issue)
    return deduped


def _material_missing_inputs(result: DealRunResult, synthesis: DealSynthesis) -> list[_MissingInputBullet]:
    items: list[_MissingInputBullet] = []
    for assessment in _critical_missing_assessments(synthesis):
        status = assessment.front_end_status or assessment.status
        if status == "conflicting across documents":
            continue
        items.append(_missing_input_from_omission(assessment))

    for item in result.issue_registry.missing_items:
        if not _missing_item_is_material(item):
            continue
        items.append(_missing_input_from_missing_item(item))

    deduped = _dedupe_missing_inputs(items)
    deduped.sort(key=lambda item: (item.priority, item.category.lower(), item.label.lower()))
    return deduped[:6]


def _missing_input_from_omission(assessment: OmissionAssessment) -> _MissingInputBullet:
    status = assessment.front_end_status or assessment.status
    request = _omission_request_text(assessment)
    return _MissingInputBullet(
        label=assessment.item,
        summary=_omission_text(assessment),
        request=request,
        category=assessment.category,
        status=status,
        priority=_MISSING_INPUT_PRIORITY.get(status, 9),
        references=_omission_reference_labels(assessment),
        document_hint=_missing_input_document_hint(
            label=assessment.item,
            category=assessment.category,
            request=request,
        ),
    )


def _missing_input_from_missing_item(item: MissingItem) -> _MissingInputBullet:
    request = item.suggested_request.rstrip(".") + "." if item.suggested_request else "Provide the current controlling support."
    priority = {
        "Purchase / Sale / Contract": 0,
        "Title": 0,
        "Vesting / Legal": 0,
        "Entitlement / Planning / Conditions": 1,
        "Utilities": 1,
        "Environmental": 1,
        "Geotech / Soils": 1,
        "Map / Plat / Improvement Plans": 2,
        "Financial / underwriting support": 2,
    }.get(item.category, 4)
    return _MissingInputBullet(
        label=item.label,
        summary=item.reason.rstrip(".") + ".",
        request=request,
        category=item.category,
        status="missing item",
        priority=priority,
        document_hint=_missing_input_document_hint(
            label=item.label,
            category=item.category,
            request=request,
        ),
    )


def _missing_item_is_material(item: MissingItem) -> bool:
    signal_text = f"{item.label} {item.category} {item.reason} {item.suggested_request}".lower()
    if _text_reads_like_extraction_noise(signal_text):
        return False
    material_categories = {
        "Purchase / Sale / Contract",
        "Title",
        "Vesting / Legal",
        "Entitlement / Planning / Conditions",
        "Environmental",
        "Geotech / Soils",
        "Utilities",
        "Map / Plat / Improvement Plans",
        "Financial / underwriting support",
    }
    if item.category in material_categories:
        return True
    return any(
        term in signal_text
        for term in (
            "purchase",
            "sale",
            "psa",
            "loi",
            "title",
            "survey",
            "vesting",
            "approval",
            "resolution",
            "utility",
            "will serve",
            "environmental",
            "geotech",
            "budget",
            "fee",
            "plan set",
        )
    )


def _dedupe_missing_inputs(items: list[_MissingInputBullet]) -> list[_MissingInputBullet]:
    seen: set[str] = set()
    deduped: list[_MissingInputBullet] = []
    for item in items:
        key = item.label.strip().lower()
        if key in seen:
            continue
        seen.add(key)
        deduped.append(item)
    return deduped


def _bottom_line_deal_text(
    result: DealRunResult,
    synthesis: DealSynthesis,
    fact_index: dict[str, list[FactRecord]],
) -> str:
    location = _location_text(fact_index)
    scale = _scale_text(fact_index)
    product = _product_text(result, fact_index)
    sentence = f"The package describes {result.deal_name} as {product}"
    if location != "jurisdiction not clearly established":
        sentence += f" in {location}"
    if scale != "scale not clearly established" and scale not in product:
        sentence += f", with {scale}"
    sentence = sentence.rstrip(".") + "."
    return clip_text(f"{sentence} Current entitlement read: {clip_text(synthesis.entitlement_status, 140)}.", 280)


def _bottom_line_issue_text(issue: CanonicalIssue) -> str:
    detail = issue.practical_impact or issue.why_it_matters or issue.likely_implication or issue.title
    return clip_text(f"{issue.title}: {detail}", 190)


def _bottom_line_missing_text(item: _MissingInputBullet) -> str:
    basis = item.request or item.summary or item.label
    return clip_text(f"{item.label}: {basis}", 190)


def _bottom_line_package_read(
    synthesis: DealSynthesis,
    material_issues: list[CanonicalIssue],
    missing_inputs: list[_MissingInputBullet],
) -> str:
    registry = synthesis.canonical_issue_registry
    read = _package_read_label(synthesis, material_issues, missing_inputs)
    quality = registry.package_quality or "mixed"
    confidence = registry.confidence_in_initial_read or "medium"
    return (
        f"Based only on the provided documents, the package currently reads as {read}. "
        f"Package quality is {quality}, and confidence in the initial read is {confidence}."
    )


def _package_read_label(
    synthesis: DealSynthesis,
    material_issues: list[CanonicalIssue],
    missing_inputs: list[_MissingInputBullet],
) -> str:
    quality = synthesis.canonical_issue_registry.package_quality
    high_severity_count = sum(issue.acquisition_severity in {"CRITICAL", "HIGH"} for issue in material_issues)
    critical_count = sum(issue.acquisition_severity == "CRITICAL" for issue in material_issues)
    if quality in {"selectively presented", "thin", "stale", "unclear"}:
        return "high-risk"
    if critical_count or high_severity_count >= 2:
        return "high-risk"
    if len(missing_inputs) >= 4:
        return "high-risk"
    if quality in {"strong", "adequate"} and not material_issues and len(missing_inputs) <= 2:
        return "relatively clean"
    return "mixed"


def _key_documents_that_control(document_analyses: list[DocumentAnalysis]) -> list[_SectionBullet]:
    ordered = sorted(
        document_analyses,
        key=lambda analysis: (
            _CONTROL_DOC_BUCKET_ORDER.get(analysis.reading_bucket, 1),
            0 if analysis.document_role == "primary" else 1,
            -analysis.reading_priority,
            analysis.document.title.lower(),
        ),
    )
    bullets: list[_SectionBullet] = []
    seen_lanes: set[str] = set()
    for analysis in ordered:
        control_reason = _control_document_reason(analysis)
        if control_reason is None:
            continue
        lane, reason = control_reason
        if lane in seen_lanes:
            continue
        seen_lanes.add(lane)
        bullets.append(_SectionBullet(text=f"{analysis.document.title}: {reason}"))
        if len(bullets) >= 6:
            break
    return bullets


def _control_document_reason(analysis: DocumentAnalysis) -> tuple[str, str] | None:
    text = f"{analysis.document.title} {analysis.document.relative_path.as_posix()}".lower()
    focus_areas = set(analysis.focus_areas)
    if any(term in text for term in ("title", "commitment", "prelim", "vesting", "deed", "exception")):
        return "title", "Controls vesting, exceptions, APN, and access/title clarity."
    if any(term in text for term in ("resolution", "approval", "conditions", "zoning", "entitlement", "development agreement")) or "Entitlement Status" in focus_areas:
        return "entitlement", "Controls approval status, open conditions, and what is actually entitled."
    if any(term in text for term in ("phase i", "phase ii", "environmental", "wetland", "biological", "habitat")) or "Environmental Risks" in focus_areas:
        return "environmental", "Controls environmental constraints and any required follow-up."
    if any(term in text for term in ("geotech", "geotechnical", "soils")) or "Geotechnical Risks" in focus_areas:
        return "geotech", "Controls soils recommendations that can move grading, retaining, and foundation cost."
    if any(term in text for term in ("will serve", "will-serve", "utility", "water", "sewer")) or "Utilities / Infrastructure Issues" in focus_areas:
        return "utilities", "Controls service availability and offsite utility obligations."
    if any(term in text for term in ("plan", "site", "grading", "improvement", "tract", "parcel map", "survey")):
        return "site", "Controls site layout, acreage, yield, and frontage assumptions."
    if any(term in text for term in ("fee", "budget", "estimate", "bid")) or focus_areas.intersection({"Fee / Exaction Burden", "Budget / Cost Reliability"}):
        return "cost", "Controls fee and cost assumptions used in underwriting."
    return None


def _critical_missing_assessments(synthesis: DealSynthesis) -> list[OmissionAssessment]:
    assessments = [
        assessment
        for assessment in synthesis.omission_assessments
        if assessment.front_end_status and assessment.front_end_status != "present and adequate"
    ]
    assessments.sort(
        key=lambda assessment: (
            _MISSING_STATUS_ORDER.get(assessment.front_end_status, 9),
            assessment.category,
            assessment.item,
        )
    )
    return assessments


def _material_conflicts(conflicts: list[ConflictRecord]) -> list[ConflictRecord]:
    return sorted(
        conflicts,
        key=lambda conflict: (_CONFLICT_TYPE_ORDER.get(conflict.fact_type, 99), conflict.label),
    )


def _build_location_bullet(fact_index: dict[str, list[FactRecord]]) -> _SectionBullet:
    fact, supporting_facts = _best_fact_bundle(fact_index, "jurisdiction")
    if fact is None:
        return _SectionBullet(text="Location / jurisdiction is not clearly established from the readable package.")
    return _SectionBullet(
        text=f"Location / jurisdiction: {fact.value}.",
        references=_reference_labels_from_facts(supporting_facts),
    )


def _build_scale_bullet(fact_index: dict[str, list[FactRecord]]) -> _SectionBullet:
    return _SectionBullet(
        text=f"Scale: {_scale_text(fact_index)}.",
        references=_scale_references(fact_index),
    )


def _build_product_bullet(result: DealRunResult, fact_index: dict[str, list[FactRecord]]) -> _SectionBullet:
    references = [
        *_reference_labels_for_fact_bundle(fact_index, "lot_count"),
        *_reference_labels_for_fact_bundle(fact_index, "unit_count"),
    ]
    return _SectionBullet(
        text=f"Product: {_product_text(result, fact_index)}.",
        references=references[:3],
    )


def _build_price_bullet(fact_index: dict[str, list[FactRecord]]) -> _SectionBullet:
    fact, supporting_facts = _best_fact_bundle(fact_index, "purchase_price")
    if fact is None:
        return _SectionBullet(text="Purchase price is not cleanly established from the readable contract support in the package.")
    qualifier = "multiple readable documents" if _support_count(fact_index.get("purchase_price", []), fact.normalized_value) >= 2 else "one clear contract document"
    return _SectionBullet(
        text=f"Purchase price: {_format_fact_sentence(fact)} This currently appears in {qualifier}.",
        references=_reference_labels_from_facts(supporting_facts),
    )


def _location_text(fact_index: dict[str, list[FactRecord]]) -> str:
    fact, _ = _best_fact_bundle(fact_index, "jurisdiction")
    return fact.value if fact is not None else "jurisdiction not clearly established"


def _scale_text(fact_index: dict[str, list[FactRecord]]) -> str:
    parts: list[str] = []
    gross, _ = _best_fact_bundle(fact_index, "gross_acreage")
    net, _ = _best_fact_bundle(fact_index, "net_acreage")
    site, _ = _best_fact_bundle(fact_index, "site_acreage")
    lots, _ = _best_fact_bundle(fact_index, "lot_count")
    units, _ = _best_fact_bundle(fact_index, "unit_count")

    if gross is not None:
        parts.append(f"gross acreage referenced at {gross.normalized_value} acres")
    if net is not None:
        parts.append(f"net acreage referenced at {net.normalized_value} acres")
    elif site is not None and gross is None:
        parts.append(f"site acreage referenced at {site.normalized_value} acres")
    if lots is not None:
        parts.append(f"{lots.normalized_value} lots")
    if units is not None:
        parts.append(f"{units.normalized_value} units")
    return ", ".join(parts) if parts else "scale not clearly established"


def _product_text(result: DealRunResult, fact_index: dict[str, list[FactRecord]]) -> str:
    lot_fact, _ = _best_fact_bundle(fact_index, "lot_count")
    unit_fact, _ = _best_fact_bundle(fact_index, "unit_count")
    document_text = " ".join(processed.document.normalized_text.lower() for processed in result.processed_documents[:12])

    for label, terms in _PRODUCT_KEYWORDS:
        if any(term in document_text for term in terms):
            if lot_fact is not None:
                return f"{label} with {lot_fact.normalized_value} lots referenced"
            if unit_fact is not None:
                return f"{label} with {unit_fact.normalized_value} units referenced"
            return label

    if lot_fact is not None:
        return f"lot-based residential subdivision with {lot_fact.normalized_value} lots referenced"
    if unit_fact is not None:
        return f"unit-based residential project with {unit_fact.normalized_value} units referenced"
    return "product type not clearly established from the readable package"


def _scale_references(fact_index: dict[str, list[FactRecord]]) -> list[str]:
    references: list[str] = []
    for fact_type in ("gross_acreage", "net_acreage", "site_acreage", "lot_count", "unit_count"):
        references.extend(_reference_labels_for_fact_bundle(fact_index, fact_type))
    return references[:3]


def _fact_section_text(fact: FactRecord, supporting_facts: list[FactRecord]) -> str:
    support_phrase = "multiple readable documents" if len(_reference_labels_from_facts(supporting_facts)) >= 2 else "one clear document"
    return f"{_format_fact_sentence(fact)} This currently appears in {support_phrase}."


def _issue_known_line(issue: CanonicalIssue) -> str:
    basis = issue.best_evidence[0] if issue.best_evidence else issue.why_it_matters or issue.likely_implication
    return clip_text(basis, 180)


def _issue_summary_line(issue: CanonicalIssue) -> str:
    summary = deal_impact_mechanism_for_issue(issue) or issue.practical_impact or issue.likely_implication or issue.why_it_matters or issue.title
    return clip_text(
        f"{issue.acquisition_severity}; {deal_impact_magnitude_for_issue(issue)} {deal_impact_type_for_issue(issue)} impact. {summary}",
        200,
    )


def _issue_section_detail(issue: CanonicalIssue) -> str:
    detail = deal_impact_mechanism_for_issue(issue) or issue.practical_impact or issue.why_it_matters or issue.likely_implication or issue.what_would_resolve_it
    return clip_text(
        f"{issue.acquisition_severity}; {deal_impact_magnitude_for_issue(issue)} {deal_impact_type_for_issue(issue)} impact. {detail}",
        200,
    )


def _issue_what_line(issue: CanonicalIssue) -> str:
    for candidate in [issue.site_specific_trigger, *issue.best_evidence[:2], *issue.core_facts[:2]]:
        if candidate and not _text_reads_like_extraction_noise(candidate):
            return clip_text(candidate, 200)
    return issue.title


def _issue_likely_explanation(issue: CanonicalIssue) -> str:
    if issue.likely_explanation:
        return clip_text(issue.likely_explanation, 220)
    if issue.status == "conflicted":
        return "Different documents appear to be using different assumptions or plan versions, and no controlling source has been established."
    if issue.status in {"not found", "unclear whether present", "present but weak"}:
        return "The package does not contain a current controlling document that cleanly closes this issue."
    category_explanations = {
        "Title / Access Concerns": "Title exceptions, easements, or access rights have not yet been reconciled to the current plan and closing structure.",
        "Entitlement Status": "Approvals may be farther along than the underlying condition closeout or supporting tracker.",
        "Environmental Risks": "Environmental follow-up remains open or not fully priced into the current underwriting assumptions.",
        "Geotechnical Risks": "Soils recommendations exist, but the current plan and budget do not clearly show they are fully carried through.",
        "Flood / Drainage Issues": "Drainage or flood-control requirements still depend on civil confirmation or permit-stage design work.",
        "Utilities / Infrastructure Issues": "Provider confirmation and offsite utility scope are still not fully locked for the current plan.",
        "Offsite Obligations": "Frontage or offsite obligations remain buyer-facing, or the cost owner is still not fixed.",
        "Fee / Exaction Burden": "Current fee support is preliminary, stale, or not fully confirmed by the governing agency.",
        "Budget / Cost Reliability": "Current cost support remains budgetary, incomplete, or not auditable enough to lock basis.",
        "Schedule Risks": "The current schedule still depends on assumptions that are not fully confirmed in the package.",
    }
    return category_explanations.get(issue.category, "The package still lacks a clean controlling basis for this issue.")


def _issue_deal_impact(issue: CanonicalIssue) -> str:
    mechanism = deal_impact_mechanism_for_issue(issue)
    if mechanism:
        return clip_text(mechanism, 220)
    if issue.practical_impact:
        return clip_text(issue.practical_impact, 220)
    impact_lines = unique_preserve_order(
        [
            issue.likely_underwriting_effect,
            issue.likely_cost_effect if issue.priority_score.cost_exposure >= 4 else "",
            issue.likely_schedule_effect if issue.priority_score.schedule_exposure >= 4 else "",
            issue.likely_closing_effect if issue.priority_score.closing_risk >= 4 else "",
            issue.likely_yield_or_product_effect if issue.priority_score.yield_exposure >= 3 else "",
            issue.why_it_matters,
        ]
    )
    filtered = [line for line in impact_lines if line]
    return clip_text(" ".join(filtered) if filtered else issue.title, 220)


def _omission_text(assessment: OmissionAssessment) -> str:
    if assessment.front_end_status == "stale and potentially unreliable":
        return f"{assessment.item} appears stale, so this lane should not be treated as current until refreshed."
    if assessment.front_end_status == "conflicting across documents":
        return f"{assessment.item} is not controlled by one current source, so this lane remains unresolved."
    return f"{assessment.item} is missing or not decision-grade in the current package."


def _missing_item_text(item: MissingItem) -> str:
    return f"{item.label}: {item.reason.rstrip('.')} Needed next: {item.suggested_request.rstrip('.')}."


def _issue_seller_question(issue: CanonicalIssue) -> str:
    return f"Please confirm {_issue_question_subject(issue)} and provide {_issue_document_hint(issue)}."


def _missing_input_seller_question(item: _MissingInputBullet) -> str:
    if item.status == "stale and potentially unreliable":
        return f"Please provide the current {item.document_hint or item.label.lower()} that controls this lane."
    return f"Please provide {item.document_hint or item.label.lower()}."


def _issue_question_subject(issue: CanonicalIssue) -> str:
    signal_text = _issue_signal_text(issue)
    if _signal_contains(signal_text, ("unit", "lot", "count", "yield", "density")):
        return "the currently approved unit and lot count"
    if _signal_contains(signal_text, ("acre", "gross", "net", "site area", "site acreage")):
        return "the controlling acreage basis"
    if issue.category == "Title / Access Concerns" or _signal_contains(signal_text, ("title", "access", "easement", "vesting", "deed", "apn", "parcel")):
        return "the current vesting, access, and title-exception position"
    if issue.category == "Entitlement Status" or _signal_contains(signal_text, ("approval", "resolution", "condition", "zoning", "permit", "map")):
        return "the current approval status, open conditions, and controlling plan"
    if issue.category == "Environmental Risks":
        return "whether any environmental follow-up, mitigation, or agency action remains open"
    if issue.category == "Geotechnical Risks":
        return "the geotechnical recommendations that still control grading, retaining, or foundation scope"
    if issue.category == "Flood / Drainage Issues":
        return "the drainage or flood-control scope that governs the current plan"
    if issue.category == "Utilities / Infrastructure Issues":
        return "current utility availability and any remaining provider conditions"
    if issue.category == "Offsite Obligations":
        return "the buyer-facing offsite or frontage scope, if any"
    if issue.category == "Fee / Exaction Burden":
        return "the current fee schedule that should control underwriting"
    if issue.category == "Budget / Cost Reliability":
        return "the current auditable site-cost basis"
    if issue.category == "Schedule Risks":
        return "the milestone schedule and the assumptions that control it"
    return f"the current position on {issue.title.lower()}"


def _issue_document_hint(issue: CanonicalIssue) -> str:
    signal_text = _issue_signal_text(issue)
    if _signal_contains(signal_text, ("unit", "lot", "count", "yield", "density")):
        return "the controlling approved plan set or resolution"
    if _signal_contains(signal_text, ("acre", "gross", "net", "site area", "site acreage")):
        return "the controlling survey, legal description, or plan sheet used for acreage"
    mapping = {
        "Title / Access Concerns": "the current title report or commitment, survey, and any exception response that controls closing",
        "Entitlement Status": "the controlling resolution, conditions tracker, or approved plan set",
        "Environmental Risks": "the current environmental report, agency correspondence, or closure documentation",
        "Geotechnical Risks": "the current geotechnical report or addendum and the plan/budget carry-through",
        "Flood / Drainage Issues": "the current drainage study, civil response, or public-works direction",
        "Utilities / Infrastructure Issues": "current will-serve letters or provider confirmation",
        "Offsite Obligations": "the current offsite scope schedule or responsibility matrix",
        "Fee / Exaction Burden": "the current city-confirmed fee schedule or fee matrix",
        "Budget / Cost Reliability": "current bid backup, estimate support, or the latest site-development budget",
        "Schedule Risks": "the current milestone schedule and the assumption backup behind it",
    }
    return mapping.get(issue.category, "the controlling document that answers the item")


def _missing_input_document_hint(*, label: str, category: str, request: str) -> str:
    signal_text = f"{label} {category} {request}".lower()
    if category == "Purchase / Sale / Contract" or _signal_contains(signal_text, ("purchase", "sale", "psa", "loi", "amendment")):
        return "the current LOI, PSA, and amendments"
    if category in {"Title", "Vesting / Legal", "Title / Access Concerns"} or _signal_contains(signal_text, ("title", "survey", "vesting", "easement", "deed")):
        return "the current title report or commitment, survey, and vesting support"
    if category in {"Entitlement / Planning / Conditions", "Entitlement Status"} or _signal_contains(signal_text, ("approval", "resolution", "condition", "zoning", "plan set")):
        return "the controlling approval set, resolution, or conditions tracker"
    if category in {"Utilities", "Utilities / Infrastructure Issues"} or _signal_contains(signal_text, ("utility", "will serve", "water", "sewer")):
        return "current will-serve letters or provider confirmation"
    if category in {"Environmental", "Environmental Risks"} or _signal_contains(signal_text, ("environmental", "phase i", "phase ii", "wetland")):
        return "the current environmental report or closure documentation"
    if category in {"Geotech / Soils", "Geotechnical Risks"} or _signal_contains(signal_text, ("geotech", "geotechnical", "soils")):
        return "the current geotechnical report or addendum"
    if category in {"Financial / underwriting support", "Budget / Cost Reliability", "Fee / Exaction Burden"} or _signal_contains(signal_text, ("budget", "bid", "pricing", "fee")):
        return "the current budget, bid backup, or fee matrix"
    if category == "Map / Plat / Improvement Plans" or _signal_contains(signal_text, ("plan", "map", "plat", "improvement")):
        return "the controlling plan set or map"
    return label.lower()


def _issue_signal_text(issue: CanonicalIssue) -> str:
    return " ".join(
        part
        for part in [
            issue.title,
            issue.site_specific_trigger,
            issue.why_it_matters,
            issue.likely_implication,
            " ".join(issue.best_evidence[:2]),
            " ".join(issue.core_facts[:2]),
        ]
        if part
    ).lower()


def _signal_contains(signal_text: str, terms: tuple[str, ...]) -> bool:
    return any(re.search(r"\b" + re.escape(term) + r"\b", signal_text) for term in terms)


def _issue_request_text(issue: CanonicalIssue) -> str:
    if issue.what_would_resolve_it:
        return issue.what_would_resolve_it.rstrip(".") + "."
    if issue.open_questions:
        return issue.open_questions[0].rstrip("?") + "?"
    return f"Please confirm how {issue.title.lower()} is being resolved and provide the controlling support."


def _issue_severity_line(issue: CanonicalIssue) -> str:
    affects = ", ".join(issue.affects[:3]) or "deal execution"
    gating = "gating item" if issue.gating_item else "non-gating item"
    return f"{issue.acquisition_severity}; {gating}; affects {affects}."


def _issue_reality_vs_noise(issue: CanonicalIssue) -> str:
    if issue.reality_vs_noise:
        return clip_text(issue.reality_vs_noise, 220)
    if issue.information_status == "conflicting across documents":
        return "Likely real inconsistency because readable documents conflict on a controlling assumption."
    if issue.information_status == "stale and potentially unreliable":
        return "Likely stale support rather than a true contradiction."
    if issue.confidence == "low":
        return "Signal remains weak and should not be over-weighted until confirmed."
    return "Likely real issue because the current readable support is specific enough to merit attention."


def _omission_request_text(assessment: OmissionAssessment) -> str:
    if assessment.recommended_request:
        request = assessment.recommended_request.rstrip(".")
        if request.lower().startswith("please "):
            return request + "."
        return f"Please provide {request}."
    return f"Please provide current, readable support for {assessment.item.lower()}."


def _reference_labels_for_fact_bundle(
    fact_index: dict[str, list[FactRecord]],
    fact_type: str,
) -> list[str]:
    fact, supporting_facts = _best_fact_bundle(fact_index, fact_type)
    if fact is None:
        return []
    return _reference_labels_from_facts(supporting_facts)


def _reference_labels_from_facts(facts: list[FactRecord]) -> list[str]:
    references: list[str] = []
    for fact in facts:
        references.extend(_source_reference_labels(fact.sources))
    return unique_preserve_order(references)[:3]


def _source_reference_labels(sources: list[SourceReference]) -> list[str]:
    labels: list[str] = []
    for source in sources:
        label = Path(source.relative_path).name
        if source.page_number is not None:
            label += f" p. {source.page_number}"
        labels.append(label)
    return unique_preserve_order(labels)[:3]


def _issue_reference_labels(issue: CanonicalIssue) -> list[str]:
    labels = [_citation_label(citation) for citation in issue.citations]
    labels.extend(issue.source_documents)
    return unique_preserve_order(labels)[:3]


def _omission_reference_labels(assessment: OmissionAssessment) -> list[str]:
    labels = [_citation_label(citation) for citation in assessment.citations]
    labels.extend(assessment.source_documents)
    return unique_preserve_order(labels)[:3]


def _citation_labels(citations: list[Citation]) -> list[str]:
    return unique_preserve_order([_citation_label(citation) for citation in citations])[:3]


def _citation_label(citation: Citation) -> str:
    label = citation.document_name
    if citation.page_number is not None:
        label += f" p. {citation.page_number}"
    return label


def _confidence_rank(confidence: str) -> int:
    return {"low": 0, "medium": 1, "high": 2}.get(confidence, 0)


def _add_section(document: Document, title: str) -> None:
    document.add_heading(title, level=1)


def _add_subsection(document: Document, title: str) -> None:
    document.add_heading(title, level=2)


def _add_bullet_items(document: Document, items: list[_SectionBullet]) -> None:
    for item in _dedupe_bullets(items):
        document.add_paragraph(item.text, style="List Bullet")
        if item.note:
            _add_note_line(document, item.note)
        if item.references:
            _add_reference_line(document, item.references)


def _add_note_line(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.space_after = Pt(1)


def _add_reference_line(document: Document, references: list[str]) -> None:
    paragraph = document.add_paragraph(f"Ref: {'; '.join(unique_preserve_order(references)[:3])}")
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.space_after = Pt(1)


def _dedupe_bullets(items: list[_SectionBullet]) -> list[_SectionBullet]:
    seen: set[str] = set()
    deduped: list[_SectionBullet] = []
    for item in items:
        key = item.text.strip().lower()
        if not key or key in seen:
            continue
        seen.add(key)
        deduped.append(item)
    return deduped


def _format_fact_sentence(fact: FactRecord) -> str:
    if fact.fact_type == "purchase_price":
        return f"Purchase price referenced at ${_format_currency(fact.normalized_value)}."
    if fact.fact_type in {"gross_acreage", "net_acreage", "site_acreage"}:
        return f"{fact.label} referenced at {fact.normalized_value} acres."
    if fact.fact_type in {"lot_count", "unit_count"}:
        count = _coerce_int(fact.normalized_value)
        count_text = str(count) if count is not None else fact.value
        noun = "lots" if fact.fact_type == "lot_count" else "units"
        return f"{fact.label} referenced at {count_text} {noun}."
    if fact.fact_type == "apn":
        return f"APN referenced as {fact.value}."
    if fact.fact_type == "owner_name":
        return f"Owner or seller referenced as {fact.value}."
    if fact.fact_type == "jurisdiction":
        return f"Jurisdiction referenced as {fact.value}."
    if fact.fact_type == "zoning":
        return f"Zoning referenced as {fact.value}."
    return fact.statement


def _format_currency(normalized_value: str) -> str:
    amount = float(normalized_value)
    return f"{amount:,.0f}" if amount.is_integer() else f"{amount:,.2f}"


def _coerce_int(value: str) -> int | None:
    try:
        return int(float(value))
    except ValueError:
        return None


def _controlling_fact_bullet(fact: AcquisitionControllingFact) -> _SectionBullet:
    rejected = f" Rejected alternatives: {', '.join(fact.rejected_alternatives[:3])}." if fact.rejected_alternatives else ""
    return _SectionBullet(
        text=f"{fact.label}: {fact.controlling_value}. Control document: {fact.controlling_document}. Why it controls: {fact.why_it_controls}{rejected}",
        references=_citation_labels(fact.citations),
    )


def _sanity_correction_bullet(item: AcquisitionSanityCorrection) -> _SectionBullet:
    return _SectionBullet(
        text=(
            f"{item.fact_type.replace('_', ' ').title()}: corrected to {item.corrected_value}. "
            f"Prior read: {item.prior_value}."
        ),
        note=f"Why prior read was wrong: {item.why_prior_was_wrong} Credible interpretation: {item.credible_interpretation}",
        references=_citation_labels(item.citations),
    )


def _acquisition_risk_bullet(item: AcquisitionRiskItem) -> _SectionBullet:
    note_parts = [
        f"Primary lever: {item.primary_lever}.",
        f"Impact: {item.impact}.",
        f"Timing: {item.timing}.",
        f"Curability: {item.curability}.",
    ]
    if item.cost_impact:
        note_parts.append(f"Cost: {item.cost_impact}")
    if item.land_value_impact:
        note_parts.append(f"Land value: {item.land_value_impact}")
    if item.margin_impact:
        note_parts.append(f"Margin: {item.margin_impact}")
    if item.irr_impact:
        note_parts.append(f"IRR: {item.irr_impact}")
    if item.timing_impact:
        note_parts.append(f"Timing impact: {item.timing_impact}")
    if any((item.price_response, item.terms_response, item.timing_response, item.contingency_response)):
        note_parts.append(
            "Structure response: "
            + " ".join(
                part
                for part in (
                    f"Price={item.price_response}" if item.price_response else "",
                    f"Terms={item.terms_response}" if item.terms_response else "",
                    f"Timing={item.timing_response}" if item.timing_response else "",
                    f"Contingency={item.contingency_response}" if item.contingency_response else "",
                )
                if part
            )
        )
    return _SectionBullet(
        text=f"{item.title}: {item.summary}",
        note=" ".join(note_parts),
        references=[*_citation_labels(item.citations), *item.source_documents],
    )


def _critical_path_bullet(step: AcquisitionCriticalPathStep) -> _SectionBullet:
    return _SectionBullet(
        text=f"Step {step.sequence}: {step.blocker}. {step.why_it_blocks}",
        references=[*_citation_labels(step.citations), *step.source_documents],
    )
